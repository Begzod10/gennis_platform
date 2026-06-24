import io
from collections import defaultdict

from django.http import HttpResponse
from django.shortcuts import get_object_or_404
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches, Cm
from rest_framework import generics, views, response, status

from group.models import Group, GroupSubjects, GroupRating
from students.models import Student
from .functions import create_multiple_years
from .models import Assignment
from .serializers import TestCreateUpdateSerializer, Test, TermSerializer, Term
from django.db.models import Case, When, Value, IntegerField, Avg, Q
from flows.models import Flow

GOLD = RGBColor(0xC9, 0xA2, 0x27)
DARK_BLUE = RGBColor(0x1A, 0x3C, 0x5E)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

ACADEMIC_YEAR = "2025-2026"

GRADE_THRESHOLDS = [
    (90, "A*"),
    (80, "A"),
    (70, "B"),
    (60, "C"),
    (50, "D"),
    (40, "E"),
    (30, "F"),
    (20, "G"),
    (0,  "U"),
]

GRADE_DESCRIPTIONS = {
    "A*": "Juda yuqori natija",
    "A":  "Excellent",
    "B":  "Very Good",
    "C":  "Good / Pass",
    "D":  "Satisfactory",
    "E":  "Minimum Pass",
    "F":  "Pass (IGCSE)",
    "G":  "Pass (IGCSE)",
    "U":  "Ungraded",
}


def _get_grade(score: float) -> str:
    for threshold, grade in GRADE_THRESHOLDS:
        if score >= threshold:
            return grade
    return "U"


def _get_level(class_number) -> str:
    if class_number is None:
        return "N/A"
    if 1 <= class_number <= 4:
        return "Primary"
    if 5 <= class_number <= 8:
        return "Secondary"
    if 9 <= class_number <= 11:
        return "Advanced"
    return "N/A"


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _add_paragraph(doc, text, size, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = 'Georgia'
    if color:
        run.font.color.rgb = color
    return p


def _build_certificate(student, term_data: list, final_score: float, grade: str, level: str) -> io.BytesIO:
    doc = Document()

    section = doc.sections[0]
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    # ── HEADER TABLE (dark blue banner) ───────────────────────────────────────
    hdr = doc.add_table(rows=1, cols=1)
    hdr.style = 'Table Grid'
    hdr_cell = hdr.rows[0].cells[0]
    _set_cell_bg(hdr_cell, '1A3C5E')
    hdr_cell.width = Inches(10.5)

    p = hdr_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("✦  CERTIFICATE OF EXCELLENCE  ✦")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'Georgia'
    run.font.color.rgb = GOLD

    doc.add_paragraph()

    # ── GOLD LINE ─────────────────────────────────────────────────────────────
    gl = doc.add_table(rows=1, cols=1)
    gl.style = 'Table Grid'
    _set_cell_bg(gl.rows[0].cells[0], 'C9A227')
    gl.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(2)
    gl.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Pt(2)

    doc.add_paragraph()

    # ── BODY ──────────────────────────────────────────────────────────────────
    _add_paragraph(doc, "This is to certify that", 13, color=DARK_BLUE, space_before=4, space_after=2)

    _add_paragraph(
        doc,
        f"{student.user.name.upper()} {student.user.surname.upper()}",
        28, bold=True, color=DARK_BLUE, space_before=4, space_after=4
    )

    _add_paragraph(doc, "has successfully completed the academic year", 13, color=DARK_BLUE, space_after=2)

    _add_paragraph(doc, ACADEMIC_YEAR, 16, bold=True, color=GOLD, space_after=10)

    _add_paragraph(doc, f"Level: {level}", 13, color=DARK_BLUE, space_after=2)

    # ── GRADE BOX ─────────────────────────────────────────────────────────────
    grade_tbl = doc.add_table(rows=1, cols=3)
    grade_tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    grade_tbl.style = 'Table Grid'
    grade_tbl.rows[0].cells[0].width = Inches(3.5)
    grade_tbl.rows[0].cells[1].width = Inches(1.5)
    grade_tbl.rows[0].cells[2].width = Inches(3.5)

    _set_cell_bg(grade_tbl.rows[0].cells[0], 'FFFFFF')
    _set_cell_bg(grade_tbl.rows[0].cells[1], '1A3C5E')
    _set_cell_bg(grade_tbl.rows[0].cells[2], 'FFFFFF')

    left_p = grade_tbl.rows[0].cells[0].paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    left_p.paragraph_format.space_before = Pt(6)
    left_r = left_p.add_run(f"Overall Score:  {final_score:.1f}")
    left_r.font.size = Pt(13)
    left_r.font.name = 'Georgia'
    left_r.font.color.rgb = DARK_BLUE

    mid_p = grade_tbl.rows[0].cells[1].paragraphs[0]
    mid_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    mid_p.paragraph_format.space_before = Pt(2)
    mid_r = mid_p.add_run(grade)
    mid_r.bold = True
    mid_r.font.size = Pt(36)
    mid_r.font.name = 'Georgia'
    mid_r.font.color.rgb = GOLD

    right_p = grade_tbl.rows[0].cells[2].paragraphs[0]
    right_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    right_p.paragraph_format.space_before = Pt(6)
    right_r = right_p.add_run(f"  {GRADE_DESCRIPTIONS.get(grade, '')}")
    right_r.font.size = Pt(13)
    right_r.font.name = 'Georgia'
    right_r.font.color.rgb = DARK_BLUE

    doc.add_paragraph()

    # ── TERM BREAKDOWN TABLE ───────────────────────────────────────────────────
    _add_paragraph(doc, "Academic Performance by Quarter", 11, bold=True, color=DARK_BLUE, space_before=6, space_after=4)

    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
    headers = ["Quarter", "Avg Score", "Grade", "Subjects"]
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        _set_cell_bg(cell, '1A3C5E')
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Georgia'
        run.font.color.rgb = WHITE

    for td in term_data:
        row = tbl.add_row()
        vals = [
            f"{td['quarter']}-chorak",
            f"{td['avg']:.1f}" if td['avg'] is not None else "—",
            td['grade'] if td['avg'] is not None else "—",
            str(td['subject_count']),
        ]
        for i, val in enumerate(vals):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(val)
            run.font.size = Pt(10)
            run.font.name = 'Georgia'
            run.font.color.rgb = DARK_BLUE

    doc.add_paragraph()

    # ── GOLD FOOTER LINE ──────────────────────────────────────────────────────
    fl = doc.add_table(rows=1, cols=1)
    fl.style = 'Table Grid'
    _set_cell_bg(fl.rows[0].cells[0], 'C9A227')
    fl.rows[0].cells[0].paragraphs[0].paragraph_format.space_before = Pt(2)
    fl.rows[0].cells[0].paragraphs[0].paragraph_format.space_after = Pt(2)

    from django.utils import timezone
    today = timezone.localdate().strftime("%B %d, %Y")
    _add_paragraph(doc, f"Issued: {today}", 9, color=DARK_BLUE, space_before=6, space_after=0)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


class CreateTest(generics.CreateAPIView):
    queryset = Test.objects.all()
    serializer_class = TestCreateUpdateSerializer


class UpdateTest(generics.UpdateAPIView):
    queryset = Test.objects.all()
    serializer_class = TestCreateUpdateSerializer


class EducationYears(generics.ListAPIView):
    queryset = Term.objects.all()
    serializer_class = TermSerializer

    def get(self, request, *args, **kwargs):
        current_year = "2025-2026"

        result = (
            Term.objects
            .values('academic_year')
            .distinct()
            .annotate(
                custom_order=Case(
                    When(academic_year=current_year, then=Value(0)),
                    default=Value(1),
                    output_field=IntegerField()
                )
            )
            .order_by('custom_order', '-academic_year')
        )

        return response.Response(result, status=status.HTTP_200_OK)


class ListTerm(generics.ListAPIView):
    queryset = Term.objects.all()
    serializer_class = TermSerializer

    def get_queryset(self):
        create_multiple_years(2025)
        academic_year = self.kwargs.get('academic_year')
        return self.queryset.filter(academic_year=academic_year)


class DeleteTest(generics.DestroyAPIView):
    queryset = Test.objects.all()
    serializer_class = TestCreateUpdateSerializer

    def delete(self, request, *args, **kwargs):
        test = self.get_object()
        test.deleted = True
        test.save()
        return response.Response({"message": "Test deleted successfully"}, status=status.HTTP_200_OK)


class ListTest(views.APIView):
    def get(self, request, *args, **kwargs):
        term = get_object_or_404(Term, id=kwargs.get('term'))
        branch_id = kwargs.get('branch')

        groups = (
            Group.objects
            .filter(deleted=False, branch_id=branch_id)
            .select_related('class_number', 'color')
            .order_by('class_number__number')
        )
        group_ids = list(groups.values_list('id', flat=True))

        flows = (
            Flow.objects
            .filter(branch_id=branch_id)
            .select_related('subject')
            .order_by('order', 'id')
        )
        flow_ids = list(flows.values_list('id', flat=True))

        gs_qs = (
            GroupSubjects.objects
            .filter(group_id__in=group_ids)
            .select_related('subject')
            .distinct('group_id', 'subject_id')
        )
        gs_by_group = defaultdict(list)
        for gs in gs_qs:
            gs_by_group[gs.group_id].append(gs)

        tests_qs = (
            Test.objects
            .filter(term=term, deleted=False)
            .filter(Q(group_id__in=group_ids) | Q(flow_id__in=flow_ids))
            .values('id', 'name', 'weight', 'date', 'group_id', 'flow_id', 'subject_id')
        )

        tests_by_group_subject = defaultdict(list)
        tests_by_flow_subject = defaultdict(list)
        for t in tests_qs:
            row = {
                "id": t['id'],
                "name": t['name'],
                "weight": t['weight'],
                "date": t['date'],
            }
            if t['group_id']:
                tests_by_group_subject[(t['group_id'], t['subject_id'])].append(row)
            elif t['flow_id']:
                tests_by_flow_subject[(t['flow_id'], t['subject_id'])].append(row)

        result = []

        for group in groups:
            class_no = getattr(group.class_number, 'number', None)
            color_name = getattr(group.color, 'name', None)
            display_name = group.name or f"{class_no or '?'} {color_name or '?'}"

            children = [
                {
                    "title": gs.subject.name,
                    "id": gs.subject.id,
                    "type": "subject",
                    "tableData": tests_by_group_subject.get((group.id, gs.subject_id), []),
                }
                for gs in gs_by_group.get(group.id, [])
            ]

            result.append({
                "title": display_name,
                "id": group.id,
                "type": "group",
                "children": children,
            })

        for flow in flows:
            subj = flow.subject
            if not subj:
                continue
            result.append({
                "title": flow.name,
                "id": flow.id,
                "type": "flow",
                "children": [{
                    "title": subj.name,
                    "id": subj.id,
                    "type": "subject",
                    "tableData": tests_by_flow_subject.get((flow.id, subj.id), []),
                }],
            })

        return response.Response(result, status=status.HTTP_200_OK)


class StudentAssignmentView(views.APIView):
    def get(self, request, *args, **kwargs):
        group_id = kwargs.get('group_id')
        test_id = kwargs.get('test_id')

        group = get_object_or_404(
            Group.objects.prefetch_related('students__user'),
            id=group_id
        )

        tests = Assignment.objects.filter(
            student__in=group.students.all(),
            test_id=test_id
        )

        test_map = {t.student_id: t for t in tests}

        result = []
        for student in group.students.all():
            assignment = test_map.get(student.id)

            if assignment:
                assignment_data = {
                    "id": assignment.id,
                    "percentage": assignment.percentage,
                    "date": assignment.date,
                    "is_editable": assignment.percentage == 0
                }
            else:
                assignment_data = {
                    "is_editable": True
                }

            result.append({
                "id": student.id,
                "name": student.user.name,
                "surname": student.user.surname,
                "assignment": assignment_data
            })

        return response.Response(result, status=status.HTTP_200_OK)


class AssignmentCreateView(views.APIView):
    def post(self, request, *args, **kwargs):
        data = request.data
        print(data)

        if not isinstance(data, list):
            return response.Response({"error": "Request body list bo'lishi kerak"}, status=status.HTTP_400_BAD_REQUEST)

        results = []
        errors = []

        for index, item in enumerate(data, start=1):
            test_id = item.get('test_id')
            percentage = item.get('percentage')
            student_id = item.get('student_id')

            if not test_id or not student_id or percentage is None:
                errors.append(
                    {"index": index, "error": "test_id, student_id va percentage maydonlari majburiy", "data": item})
                continue

            assignment, created = Assignment.objects.update_or_create(test_id=test_id, student_id=student_id,
                                                                      defaults={"percentage": percentage, })

            results.append({"student_id": student_id, "assignment_id": assignment.id, "created": created,
                            "result": percentage * assignment.test.weight / 100})

        response_data = {"success": results, "errors": errors}

        return response.Response(response_data, status=status.HTTP_200_OK)


class TermsByGroup(views.APIView):
    def get(self, request, *args, **kwargs):
        group_id = kwargs.get('group_id')
        term_id = kwargs.get('term_id')

        group = Group.objects.get(id=group_id)
        students = group.students.all()

        response_data = []

        for student in students:
            subject_id = kwargs.get('subject_id', None)
            if subject_id:
                assignments = Assignment.objects.filter(student=student, test__term_id=term_id,
                                                        test__subject_id=subject_id).select_related(
                    'test__subject')
            else:
                assignments = Assignment.objects.filter(student=student, test__term_id=term_id).select_related(
                    'test__subject')

            subjects_data = defaultdict(
                lambda: {"subject_name": "", "assignments": [], "total_result": 0, "count": 0, "average_result": 0})

            for assignment in assignments:
                subject_id = assignment.test.subject.id
                calculated_result = (assignment.test.weight * assignment.percentage) / 100

                subjects_data[subject_id]["subject_name"] = assignment.test.subject.name
                subjects_data[subject_id]["assignments"].append(
                    {"test_name": assignment.test.name, "percentage": assignment.percentage,
                     "calculated_result": round(calculated_result, 2)})
                subjects_data[subject_id]["total_result"] += calculated_result
                subjects_data[subject_id]["count"] += 1

            for subj_id, subj_data in subjects_data.items():
                if subj_data["count"] > 0:
                    subj_data["average_result"] =subj_data["total_result"]
                del subj_data["total_result"]
                del subj_data["count"]

            response_data.append({"id": student.id, "first_name": student.user.name, "last_name": student.user.surname,
                                  "subjects": list(subjects_data.values())})

        return response.Response(response_data, status=status.HTTP_200_OK)


class TermsByStudent(views.APIView):
    def get(self, request, *args, **kwargs):
        student_id = kwargs.get('student_id')
        term_id = kwargs.get('term_id')
        subject_id = kwargs.get('subject_id', None)

        student = Student.objects.get(id=student_id)
        if subject_id:

            assignments = Assignment.objects.filter(
                student=student,
                test__term_id=term_id,
                test__subject_id=subject_id
            ).select_related('test__subject')
        else:
            assignments = Assignment.objects.filter(
                student=student,
                test__term_id=term_id
            ).select_related('test__subject')

        response_data = {
            "student": {
                "first_name": student.user.first_name,
                "last_name": student.user.last_name,
            },
            "subjects": [],
            "total_result": 0,
            "average_result": 0
        }

        subjects_data = defaultdict(lambda: {
            "subject_name": "",
            "assignments": [],
            "total_result": 0,
            "count": 0,
            "average_result": 0
        })

        total_result_all = 0
        count_all = 0

        for assignment in assignments:
            subject_id = assignment.test.subject.id
            calculated_result = (assignment.test.weight * assignment.percentage) / 100

            subjects_data[subject_id]["subject_name"] = assignment.test.subject.name
            subjects_data[subject_id]["assignments"].append({
                "test_name": assignment.test.name,
                "percentage": assignment.percentage,
                "calculated_result": round(calculated_result, 2)
            })
            subjects_data[subject_id]["total_result"] += calculated_result
            subjects_data[subject_id]["count"] += 1

            total_result_all += calculated_result
            count_all += 1

        for subject_id, subj_data in subjects_data.items():
            if subj_data["count"] > 0:
                subj_data["average_result"] = subj_data["total_result"]
            del subj_data["count"]
            del subj_data["total_result"]
            response_data["subjects"].append(subj_data)

        if count_all > 0:
            response_data["total_result"] = round(total_result_all, 2)
            response_data["average_result"] =total_result_all

        return response.Response(response_data, status=status.HTTP_200_OK)


class GroupSubjectsApiView(views.APIView):
    def get(self, request, *args, **kwargs):
        group_id = kwargs.get('group_id')
        group = Group.objects.get(id=group_id)

        subjects = GroupSubjects.objects.filter(group=group).select_related('subject').values(
            'subject__id', 'subject__name'
        )

        subjects_data = [
            {"id": subj["subject__id"], "name": subj["subject__name"]} for subj in subjects
        ]

        return response.Response(subjects_data, status=status.HTTP_200_OK)

class EducationQualityOverview(views.APIView):
    def get(self, request, *args, **kwargs):
        def get_param(names):
            for name in names:
                val = request.query_params.get(name)
                if val and val not in ['undefined', 'null', '', 'all']:
                    return val
            return None

        branch_id = get_param(['branch_id', 'branch'])
        subject_id = get_param(['subject_id', 'subject'])
        class_id = get_param(['class_id', 'class', 'group_id'])
        teacher_id = get_param(['teacher_id', 'teacher'])
        term_id = get_param(['term_id', 'term'])

        # Include group-tests (non-deleted group) and flow-tests
        assignments = Assignment.objects.filter(test__deleted=False).filter(
            Q(test__group__deleted=False) | Q(test__flow__isnull=False)
        )

        if branch_id:
            assignments = assignments.filter(
                Q(test__group__branch_id=branch_id) | Q(test__flow__branch_id=branch_id)
            )
        if subject_id:
            assignments = assignments.filter(test__subject_id=subject_id)
        if class_id:
            assignments = assignments.filter(test__group_id=class_id)
        if teacher_id:
            assignments = assignments.filter(
                Q(test__group__teacher=teacher_id) | Q(test__flow__teacher_id=teacher_id)
            )
        if term_id:
            assignments = assignments.filter(test__term_id=term_id)

        avg_percentage = assignments.aggregate(Avg('percentage'))['percentage__avg'] or 0
        avg_rating = avg_percentage / 20
        return response.Response({
            "rating": round(avg_rating, 1),
            "max_rating": 5,
            "description": "5 ballik tizimda maktab reytingi"
        }, status=status.HTTP_200_OK)


class StudentCertificateView(views.APIView):
    """
    GET /api/terms/certificate/<student_id>/

    2025-2026 o'quv yili uchun sertifikat generatsiya qiladi.
    - Har chorak uchun: fanlar ball yig'indisi / fanlar soni = chorak o'rtachasi
    - Yakuniy ball: 4 chorak o'rtachasi / ma'lumot bor choraklar soni
    - Daraja: class_number bo'yicha (1-4 Primary, 5-8 Secondary, 9-11 Advanced)
    """

    def get(self, request, student_id):
        student = get_object_or_404(
            Student.objects.select_related('user', 'class_number'),
            id=student_id
        )

        terms = list(
            Term.objects.filter(academic_year=ACADEMIC_YEAR).order_by('quarter')
        )

        term_data = []
        total_avg = 0.0
        counted_terms = 0

        for term in terms:
            assignments = (
                Assignment.objects
                .filter(student=student, test__term=term, test__deleted=False)
                .select_related('test__subject')
            )

            subject_scores = defaultdict(float)
            for a in assignments:
                score = (a.test.weight * a.percentage) / 100
                subject_scores[a.test.subject_id] += score

            subject_count = len(subject_scores)
            if subject_count > 0:
                term_avg = sum(subject_scores.values()) / subject_count
                total_avg += term_avg
                counted_terms += 1
            else:
                term_avg = None

            term_data.append({
                'quarter': term.quarter,
                'avg': term_avg,
                'grade': _get_grade(term_avg) if term_avg is not None else None,
                'subject_count': subject_count,
            })

        if counted_terms == 0:
            return response.Response(
                {"detail": "Bu o'quvchi uchun 2025-2026 o'quv yilida hech qanday baho topilmadi."},
                status=status.HTTP_404_NOT_FOUND
            )

        final_score = total_avg / counted_terms
        grade = _get_grade(final_score)
        level = _get_level(
            student.class_number.number if student.class_number else None
        )

        buf = _build_certificate(student, term_data, final_score, grade, level)

        name = f"{student.user.name}_{student.user.surname}".replace(' ', '_')
        resp = HttpResponse(
            buf.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        resp['Content-Disposition'] = f'attachment; filename="certificate_{name}_{ACADEMIC_YEAR}.docx"'
        return resp


class EducationQualityDetails(views.APIView):
    def get(self, request, *args, **kwargs):
        term_id = kwargs.get('term_id')
        term = get_object_or_404(Term, id=term_id)

        def get_param(names):
            for name in names:
                val = request.query_params.get(name)
                if val and val not in ['undefined', 'null', '', 'all']:
                    return val
            return None

        subject_id = get_param(['subject_id', 'subject'])
        class_id = get_param(['class_id', 'class', 'group_id'])
        teacher_id = get_param(['teacher_id', 'teacher'])
        branch_id = get_param(['branch_id', 'branch'])

        assignments = Assignment.objects.filter(
            test__term_id=term_id, test__deleted=False
        ).filter(
            Q(test__group__deleted=False) | Q(test__flow__isnull=False)
        )

        if branch_id:
            assignments = assignments.filter(
                Q(test__group__branch_id=branch_id) | Q(test__flow__branch_id=branch_id)
            )
        if teacher_id:
            assignments = assignments.filter(
                Q(test__group__teacher=teacher_id) | Q(test__flow__teacher_id=teacher_id)
            )
        if class_id:
            assignments = assignments.filter(test__group_id=class_id)
        if subject_id:
            assignments = assignments.filter(test__subject_id=subject_id)

        chart_data = []
        chart_type = "subjects"

        if subject_id:
            chart_type = "classes"
            group_data = (
                assignments.filter(test__group__isnull=False)
                .values('test__group__name', 'test__group__class_number__number')
                .annotate(avg=Avg('percentage'))
                .order_by('test__group__class_number__number')
            )
            for item in group_data:
                label = item['test__group__name'] or f"{item['test__group__class_number__number']}-sinf"
                chart_data.append({"label": label, "value": round((item['avg'] or 0) / 20, 1)})

            flow_data = (
                assignments.filter(test__flow__isnull=False)
                .values('test__flow__name')
                .annotate(avg=Avg('percentage'))
                .order_by('test__flow__name')
            )
            for item in flow_data:
                chart_data.append({
                    "label": item['test__flow__name'],
                    "value": round((item['avg'] or 0) / 20, 1),
                })

        elif teacher_id:
            chart_type = "performance"
            group_data = (
                assignments.filter(test__group__isnull=False)
                .values('test__group__name', 'test__subject__name')
                .annotate(avg=Avg('percentage'))
            )
            for item in group_data:
                label = f"{item['test__subject__name']} ({item['test__group__name']})"
                chart_data.append({"label": label, "value": round((item['avg'] or 0) / 20, 1)})

            flow_data = (
                assignments.filter(test__flow__isnull=False)
                .values('test__flow__name', 'test__subject__name')
                .annotate(avg=Avg('percentage'))
            )
            for item in flow_data:
                label = f"{item['test__subject__name']} ({item['test__flow__name']})"
                chart_data.append({"label": label, "value": round((item['avg'] or 0) / 20, 1)})

        elif class_id:
            chart_type = "subjects"
            data = (
                assignments.values('test__subject__name')
                .annotate(avg=Avg('percentage'))
            )
            for item in data:
                chart_data.append({"label": item['test__subject__name'], "value": round((item['avg'] or 0) / 20, 1)})

        else:
            data = (
                assignments.values('test__subject__name')
                .annotate(avg=Avg('percentage'))
                .order_by('-avg')
            )
            for item in data:
                if item['test__subject__name']:
                    chart_data.append({"label": item['test__subject__name'], "value": round((item['avg'] or 0) / 20, 1)})

        return response.Response({
            "term": TermSerializer(term).data,
            "charts": [
                {
                    "type": chart_type,
                    "data": chart_data
                }
            ]
        }, status=status.HTTP_200_OK)
