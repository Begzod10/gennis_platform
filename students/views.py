import calendar
import datetime
import json
import os
import uuid
from datetime import date
from datetime import datetime

import docx
from django.db.models import Prefetch
from django.db.models import Q
from django.db.models import Sum
from django.db.models.functions import ExtractMonth
from django.db.models.functions import ExtractYear
from django.shortcuts import get_object_or_404
from django.utils import timezone
from rest_framework import filters
from rest_framework import generics
from rest_framework import status
from rest_framework.generics import ListAPIView
from rest_framework.permissions import IsAuthenticated
from rest_framework.response import Response
from rest_framework.views import APIView

from attendances.models import AttendancePerMonth
from branch.models import Branch
from classes.models import ClassNumber
from group.models import Group
from permissions.response import QueryParamFilterMixin
from students.serializer.lists import ActiveListSerializer, ActiveListDeletedStudentSerializer
from .models import Student, DeletedStudent, ContractStudent, DeletedNewStudent, StudentPayment
from .serializers import StudentCharity, get_remaining_debt_for_student
from .serializers import (StudentListSerializer, DeletedNewStudentListSerializer, StudentPaymentListSerializer)


class StudentListView(ListAPIView):
    filter_mappings = {'branch': 'user__branch_id', }
    queryset = Student.objects.all()
    serializer_class = StudentListSerializer
    filter_backends = [filters.SearchFilter]
    search_fields = ['user__name', 'user__surname', 'user__username']


class DeletedFromRegistered(QueryParamFilterMixin, ListAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = DeletedNewStudentListSerializer
    filter_mappings = {'branch': 'student__user__branch_id', 'subject': 'subject__id',
                       'age': 'student__user__birth_date', 'language': 'student__user__language_id', }

    def get_queryset(self):
        return DeletedNewStudent.objects.all()


class DeletedGroupStudents(QueryParamFilterMixin, ListAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = ActiveListDeletedStudentSerializer
    filter_mappings = {'branch': 'student__user__branch_id', 'subject': 'student__subject__id',
                       'age': 'student__user__birth_date', 'language': 'student__user__language_id', }
    filter_backends = [filters.SearchFilter]
    search_fields = ['student__user__name', 'student__user__surname', 'student__user__username']

    def get_queryset(self):
        deleted_new_student_ids = DeletedNewStudent.objects.values_list('student_id', flat=True)
        deleted = DeletedStudent.objects.filter(deleted=False).values_list('student_id', flat=True)

        queryset = DeletedStudent.objects.filter(student__id__in=deleted, deleted=False).exclude(
            student__id__in=deleted_new_student_ids).order_by('-deleted_date')

        return queryset


# @method_decorator(cache_page(60 * 2), name='dispatch')s
class NewRegisteredStudents(QueryParamFilterMixin, ListAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = ActiveListSerializer

    filter_mappings = {'branch': 'user__branch_id', 'subject': 'subject__id', 'age': 'user__birth_date',
                       'language': 'user__language_id', 'number': 'class_number_id', }

    filter_backends = [filters.SearchFilter]
    search_fields = ['user__name', 'user__surname', 'user__username']

    def get_queryset(self):
        excluded_ids = list(DeletedStudent.objects.filter(deleted=False).values_list('student_id', flat=True)) + list(
            DeletedNewStudent.objects.values_list('student_id', flat=True))

        return (
            Student.objects.filter(~Q(id__in=excluded_ids) & Q(groups_student__isnull=True)).distinct().order_by('-pk')
            # or 'id'
        )


class ActiveStudents(QueryParamFilterMixin, ListAPIView):
    permission_classes = [IsAuthenticated]
    serializer_class = ActiveListSerializer
    filter_mappings = {'branch': 'user__branch_id', 'subject': 'subject__id', 'age': 'user__birth_date',
                       'language': 'user__language_id', }
    filter_backends = [filters.SearchFilter]
    search_fields = ['user__name', 'user__surname', 'user__username']

    def get_queryset(self, *args, **kwargs):
        # with silk_profile(name='ActiveStudents.get_queryset'):
        deleted_student_ids = DeletedStudent.objects.filter(student__groups_student__isnull=True,
                                                            deleted=False).values_list('student_id', flat=True)

        deleted_new_student_ids = DeletedNewStudent.objects.values_list('student_id', flat=True)

        active_students = Student.objects.select_related('user', 'user__language', 'class_number').prefetch_related(
            'user__student_user',
            Prefetch('groups_student', queryset=Group.objects.select_related('class_number', 'color').order_by('id'),
                     to_attr='prefetched_groups')).exclude(id__in=deleted_student_ids).exclude(
            id__in=deleted_new_student_ids).filter(groups_student__isnull=False).distinct().order_by(
            'class_number__number')

        return active_students


class CreateContractView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, user_id):
        data = request.data

        calendar_year, calendar_month, calendar_day = self._get_calendar_date()
        calendar_year = str(calendar_year)  # Ensure calendar_year is a string
        student = get_object_or_404(Student, id=user_id)

        try:
            ot = datetime.strptime(data['date']['ot'], "%Y-%m-%d").date()
            do = datetime.strptime(data['date']['do'], "%Y-%m-%d").date()
        except KeyError as e:
            return Response({"error": f"Sanani aniqlashda xato: {e}"}, status=status.HTTP_400_BAD_REQUEST)
        except ValueError:
            return Response({"error": "Sanalar noto'g'ri formatda, YYYY-MM-DD formatida kiriting"},
                            status=status.HTTP_400_BAD_REQUEST)

        student.representative_name = data.get('name', student.representative_name)
        student.representative_surname = data.get('surname', student.representative_surname)
        student.save()

        ot_month, do_month = ot.month, do.month
        do_year = do.year

        if do_year > int(calendar_year):
            do_month += 12

        month = do_month - ot_month + 1
        user = student.user
        location = get_object_or_404(Branch, id=user.branch_id)
        contract = ContractStudent.objects.filter(student_id=student.id).first()
        student_charity = StudentCharity.objects.filter(student_id=student.id)
        all_charity = sum(char.charity_sum for char in student_charity)

        if contract and contract.contract and os.path.exists(contract.contract.path):
            os.remove(contract.contract.path)
            contract.contract = ""

        contract_data = ContractStudent.objects.filter(student_id=student.id, year=datetime.now()).first()

        if not contract:
            contract = ContractStudent(student=student, created_date=ot, expire_date=do,
                                       father_name=data.get('fatherName', ''), given_place=data.get('givenPlace', ''),
                                       place=data.get('place', ''), passport_series=data.get('passportSeries', ''),
                                       given_time=data.get('givenTime', ''), year=datetime.now())
            contract.save()
            if not contract_data:
                ContractStudent(year=datetime.now(), student=student).save()
            else:
                contract_data.number += 1
                contract_data.save()
        else:
            ContractStudent.objects.filter(id=contract.id).update(created_date=ot, expire_date=do,
                                                                  father_name=data.get('fatherName',
                                                                                       contract.father_name),
                                                                  given_place=data.get('givenPlace',
                                                                                       contract.given_place),
                                                                  place=data.get('place', contract.place),
                                                                  passport_series=data.get('passportSeries',
                                                                                           contract.passport_series),
                                                                  given_time=data.get('givenTime', contract.given_time))

        doc = docx.Document('media/contract.docx')
        user_name, user_surname, father_name = (
            (user.name, user.surname, user.father_name) if int(user.calculate_age()) >= 18 else (
                student.representative_name, student.representative_surname, contract.father_name))

        id_hex = uuid.uuid1().hex[:15]
        campus_name = f"{location.name} {location.name}" if location.location_type == "Shahri" else f"{location.district} {location.location_type}"
        number = f'{calendar_year}/{location.code}/{1}'

        self._fill_doc(doc, location, contract, user, student, id_hex, campus_name, number, ot, do, month, all_charity)

        doc_path = f"media/contracts/{id_hex} {user.name.title()} {user.surname.title()}doc.docx"
        doc.save(doc_path)
        contract.contract = doc_path
        contract.save()

        return Response({"success": True, "msg": "Shartnoma muvaffaqiyatli yaratildi", "file": doc_path})

    def _get_calendar_date(self):
        now = datetime.now()
        return now.year, now.month, now.day

    def _fill_doc(self, doc, location, contract, user, student, id_hex, campus_name, number, ot, do, month,
                  all_charity):
        doc.paragraphs[0].runs[0].text = f"SHARTNOMA № {number}"
        doc.paragraphs[3].text = f"{campus_name} {ot.strftime('%d-%m-%Y')}"
        if location.id > 3:
            doc.paragraphs[
                5].text = "O‘zbekiston Respublikasi Prezidentining 15.09.2017-yildagi PQ-3276 sonli “Nodavlat taʼlim xizmatlarini yanada rivojlantirish chora-tadbirlari toʻgʻrisida”gi qaroriga."
        else:
            doc.paragraphs[5].text = ""

        doc.paragraphs[
            6].text = f"Taʼlim muassasasi, yaʼni {location.campus_name} (keyingi o‘rinlarda “Nodavlat taʼlim muassasasi” deb yuritiladi), direktor {location.director_fio} nomidan va {student.representative_surname.title()} {student.representative_name.title()} {contract.father_name[0].title()}{contract.father_name[1:].lower()} bilan bir tomondan"
        doc.paragraphs[
            9].text = f"1.1 Ushbu shartnomaga muvofiq, o‘quvchining ota-onasi (yoki qonuniy vakili) o‘zining voyaga yetmagan farzandini {user.name.title()} {user.surname.title()} {user.father_name[0].title()}{user.father_name[1:].lower()} ni qo‘shimcha taʼlim olish uchun qabul qilmoqda."
        doc.paragraphs[
            15].text = f"2.1 O‘quvchining nodavlat taʼlim muassasida taʼlim olishi uchun bir oylik to‘lov summasi {abs(student.total_payment_month) - all_charity} va {contract.expire_date.strftime('%d-%m-%Y')} muddatgacha {abs(((student.total_payment_month) - all_charity) * month)} so‘mni tashkil etadi."
        doc.paragraphs[
            69].text = f"7.1 Ushbu shartnoma tomonlar o‘rtasida imzolangan kundan boshlab yuridik kuchga ega bo‘ladi va {contract.expire_date.strftime('%d-%m-%Y')} muddatga qadar amal qiladi."

        info = [{"left_info": f"{location.campus_name} NTM",
                 "right_info": f"F.I.O: {student.representative_surname.title()} {student.representative_name.title()} {contract.father_name[0].title()}{contract.father_name[1:].lower()}"},
                {"left_info": location.address,
                 "right_info": f"Pasport maʼlumotlari: Seriya {contract.passport_series}"},
                {"left_info": f"R/S: {location.bank_sheet}  INN: {location.inn}",
                 "right_info": f"Berilgan vaqti: {contract.given_time}"},
                {"left_info": f"Bank: {location.bank}", "right_info": f"Manzil: {contract.place}"},
                {"left_info": f"MFO: {location.mfo}", "right_info": ""},
                {"left_info": f"Tel: {location.number}", "right_info": ""},
                {"left_info": f"Direktor: __________{location.director_fio}", "right_info": ""},
                {"left_info": "M.P", "right_info": "Imzo____________"}]

        table = doc.add_table(rows=1, cols=2)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Nodavlat taʼlim muassasasi'
        hdr_cells[1].text = 'O‘quvchining qonuniy vakili (ota-onasi)'

        for item in info:
            row_cells = table.add_row().cells
            row_cells[0].text = item['left_info']
            row_cells[1].text = item['right_info']


class UploadPDFContractView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, user_id):
        student = get_object_or_404(ContractStudent, student__id=user_id)
        file = request.FILES.get('file')
        student.contract = file
        student.save()

        return Response({"success": True, "msg": "File uploaded successfully", "url": str(student.contract)})


class PaymentDatas(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request, student_id):
        data = json.loads(request.body)
        year = data['year']
        month = data['month']
        payments = StudentPayment.objects.filter(deleted=False, student_id=student_id, date__year=year,
                                                 date__month=month).all()
        data = StudentPaymentListSerializer(payments, many=True).data

        return Response(data)

    def get(self, request, student_id):

        student_payments = StudentPayment.objects.filter(deleted=False, student_id=student_id).all()

        payments_by_year = {}

        for payment in student_payments:
            year = payment.date.year
            month = payment.date.month

            if year not in payments_by_year:
                payments_by_year[year] = set()
            payments_by_year[year].add(month)

        payments_by_year_list = [{'name': year, 'months': sorted(months)} for year, months in payments_by_year.items()]

        return Response({'payments_by_year': payments_by_year_list, })


class GetMonth(APIView):

    def get(self, request, student_id, attendance_id):
        student = Student.objects.get(pk=student_id)
        group = student.groups_student.first()
        if group is None:
            return Response([])
        month = AttendancePerMonth.objects.filter(student_id=student_id, status=False,
                                                  group_id=group.id).all().order_by('month_date__year',
                                                                                    'month_date__month')
        data = []
        for mont in month:
            if isinstance(mont.month_date, str):
                month_date = datetime.strptime(mont.month_date, "%Y-%m-%d")
            else:
                month_date = mont.month_date
            month_name = month_date.strftime("%B")
            month_number = month_date.strftime("%m")
            data.append({'id': mont.id, 'name': month_name, 'number': month_number, 'price': mont.remaining_debt})

        return Response(data)

    def post(self, request, student_id, attendance_id):

        from attendances.models import AttendancePerMonth
        attendance_id = request.data.get('id')

        attendance_per_month = AttendancePerMonth.objects.get(pk=attendance_id)

        if attendance_per_month.total_debt != attendance_per_month.payment and attendance_per_month.remaining_debt == 0:
            attendance_per_month.remaining_debt = attendance_per_month.total_debt
            attendance_per_month.save()

        payment_sum = int(request.data['payment_sum'])
        branch = request.data['branch']

        if attendance_per_month.remaining_debt >= payment_sum:
            attendance_per_month.remaining_debt -= payment_sum
            attendance_per_month.payment += payment_sum
            student_payment = StudentPayment.objects.create(student_id=student_id, payment_sum=payment_sum,
                                                            branch_id=branch, status=request.data['status'],
                                                            payment_type_id=request.data['payment_type'],
                                                            date=request.data['date'], attendance=attendance_per_month)
            student_payment.save()

            if attendance_per_month.remaining_debt == 0:
                attendance_per_month.status = True

        attendance_per_month.save()

        from attendances.serializers import AttendancePerMonthSerializer
        get_remaining_debt_for_student(student_id)

        return Response(AttendancePerMonthSerializer(attendance_per_month).data)


class shahakota(APIView):
    def post(self, request):
        from attendances.models import AttendancePerMonth
        attendance_id = request.data['id']
        month = AttendancePerMonth.objects.get(id=attendance_id, status=False)
        if month.remaining_debt == 0 and month.payment == 0:

            data = {'price': month.total_debt}
        else:
            data = {'price': month.remaining_debt}
        return Response(data)


class DeleteStudentPayment(APIView):
    def delete(self, request, pk):
        student_payment = get_object_or_404(StudentPayment, id=pk)
        attendance_per_month = get_object_or_404(AttendancePerMonth, id=student_payment.attendance.id)

        attendance_per_month.remaining_debt += student_payment.payment_sum
        attendance_per_month.payment -= student_payment.payment_sum
        student_payment.deleted = True
        student_payment.save()

        if attendance_per_month.remaining_debt != 0:
            attendance_per_month.status = False

        attendance_per_month.save()

        return Response({'msg': "Success"}, status=status.HTTP_200_OK)


class DeleteFromDeleted(APIView):
    def delete(self, request, pk):
        deleted = DeletedStudent.objects.filter(student_id=pk).all()
        for i in deleted:
            i.deleted = True
            i.save()
        return Response({'msg': "Student muvoffaqiyatlik orqaga qaytarildi"}, status=status.HTTP_200_OK)


class MissingAttendanceListView(generics.RetrieveAPIView):

    def get_queryset(self):
        student_id = self.kwargs.get('student_id')
        student = Student.objects.get(id=student_id)
        group = student.groups_student.first()
        today = timezone.localdate()
        academic_start_year = today.year if today.month >= 9 else (today.year - 1)

        # Optional explicit academic start: ?ay=YYYY
        ay_param = self.request.query_params.get("ay")
        if ay_param and ay_param.isdigit():
            academic_start_year = int(ay_param)

        start_date = date(academic_start_year, 9, 1)  # inclusive
        end_date = date(academic_start_year + 1, 7, 1)  # exclusive
        print(start_date, end_date)
        if not group:
            deleted_student = DeletedStudent.objects.filter(
                student_id=student_id
            ).order_by("-pk").first()
            print('deleted_group', deleted_student.group_id)
            if deleted_student:
                qs = (
                    AttendancePerMonth.objects.filter(
                        student_id=student_id,
                        group_id=deleted_student.group_id,
                        month_date__gte=start_date,
                        month_date__lt=end_date,
                        group__deleted=False
                    ).annotate(
                        month_number=ExtractMonth('month_date'),
                        year_number=ExtractYear('month_date'),
                    ).order_by('month_date')
                )
                return qs

        else:
            qs = (
                AttendancePerMonth.objects.filter(student_id=student_id, group_id=group.id, month_date__gte=start_date,
                                                  month_date__lt=end_date, ).annotate(
                    month_number=ExtractMonth('month_date'),
                    year_number=ExtractYear('month_date'), ).order_by('month_date'))

            return qs

    def retrieve(self, request, *args, **kwargs):
        queryset = self.get_queryset()
        if not queryset.exists():
            return Response([])
        student_id = self.kwargs.get('student_id')

        months_with_attendance = queryset.values_list('month_number', flat=True).distinct()
        data = []
        student = Student.objects.get(pk=student_id)

        group = student.groups_student.first()
        if not group:
            deleted_student = DeletedStudent.objects.filter(student_id=student_id).order_by("-id").first()
            group = deleted_student.group
        attendances = AttendancePerMonth.objects.filter(student_id=student_id, group_id=group.id).all().order_by(
            'month_date__year', 'month_date__month')

        for attendance in attendances:
            student_payemnt = StudentPayment.objects.filter(attendance=attendance, status=True

                                                            ).first()
            data.append({'id': attendance.id, 'month': attendance.month_date, 'total_debt': attendance.total_debt,
                         'remaining_debt': attendance.remaining_debt, 'payment': attendance.payment,
                         "discount": attendance.discount, "old_money": attendance.old_money,
                         "discount_sum": student_payemnt.payment_sum if student_payemnt else 0,
                         "discount_reason": student_payemnt.reason if student_payemnt else 0,
                         "discount_id": student_payemnt.id if student_payemnt else 0,
                         "reason": StudentCharity.objects.filter(
                             student_id=student_id).first().name if StudentCharity.objects.filter(
                             student_id=student_id).first() else None, 'cash':
                             attendance.studentpayment_set.filter(payment_type__name='cash', deleted=False,
                                                                  status=False,

                                                                  ).aggregate(total_sum=Sum('payment_sum'))[
                                 'total_sum'] or 0, 'bank':
                             attendance.studentpayment_set.filter(payment_type__name='bank', deleted=False,
                                                                  status=False,

                                                                  ).aggregate(total_sum=Sum('payment_sum'))[
                                 'total_sum'] or 0, 'click':
                             attendance.studentpayment_set.filter(payment_type__name='click', deleted=False,
                                                                  status=False,

                                                                  ).aggregate(total_sum=Sum('payment_sum'))[
                                 'total_sum'] or 0,

                         })

        all_months = [9, 10, 11, 12, 1, 2, 3, 4, 5, 6]
        missing_months = set(all_months) - set(months_with_attendance)
        month_names = [calendar.month_name[month] for month in sorted(missing_months)]

        return Response({"month": month_names, "data": data})


class MissingAttendanceView(APIView):
    def post(self, request, student_id):
        student = Student.objects.get(pk=student_id)
        group = student.groups_student.first()
        if group is None:
            return Response([])
        data = json.loads(request.body)
        month = data['month']
        year = data['year']
        old_months = ["September", "October", "November", "December"]
        # if month in old_months:
        #     year = datetime.now().year - 1
        # else:
        #     year = datetime.now().year

        month_date = datetime.strptime(f"01 {month} {year}", "%d %B %Y")
        attendance = AttendancePerMonth.objects.create(student_id=student_id, month_date=month_date, group_id=group.id,
                                                       total_debt=group.price if group.price else group.class_number.price,
                                                       system=group.system

                                                       )
        attendances = AttendancePerMonth.objects.filter(student_id=student_id, group_id=group.id).all().order_by(
            'month_date__year', 'month_date__month')
        queryset = AttendancePerMonth.objects.filter(student_id=student_id,
                                                     month_date__month__in=[9, 10, 11, 12, 1, 2, 3, 4, 5, 6]).annotate(
            month_number=ExtractMonth('month_date'))
        months_with_attendance = queryset.values_list('month_number', flat=True).distinct()
        all_months = [9, 10, 11, 12, 1, 2, 3, 4, 5, 6]
        missing_months = set(all_months) - set(months_with_attendance)
        month_names = [calendar.month_name[month] for month in sorted(missing_months)]
        student_payemnt = StudentPayment.objects.filter(attendance=attendance, status=True

                                                        ).first()
        data = []
        for attendance in attendances:
            data.append({'id': attendance.id, 'month': attendance.month_date, 'total_debt': attendance.total_debt,
                         'remaining_debt': attendance.remaining_debt, "discount": attendance.discount,
                         "discount_sum": student_payemnt.payment_sum if student_payemnt else 0,
                         "discount_reason": student_payemnt.reason if student_payemnt else 0,
                         "discount_id": student_payemnt.id if student_payemnt else 0,
                         "reason": StudentCharity.objects.filter(
                             student_id=student_id).first().name if StudentCharity.objects.filter(
                             student_id=student_id).first() else None, 'cash':
                             attendance.studentpayment_set.filter(payment_type__name='cash', deleted=False,
                                                                  status=False,

                                                                  ).aggregate(total_sum=Sum('payment_sum'))[
                                 'total_sum'] or 0, 'bank':
                             attendance.studentpayment_set.filter(payment_type__name='bank', deleted=False,
                                                                  status=False,

                                                                  ).aggregate(total_sum=Sum('payment_sum'))[
                                 'total_sum'] or 0, 'click':
                             attendance.studentpayment_set.filter(payment_type__name='click', deleted=False,
                                                                  status=False, ).aggregate(
                                 total_sum=Sum('payment_sum'))['total_sum'] or 0,

                         })
        return Response({"month": month_names, "data": data, "msg": "Qarz muvaffaqiyatli yaratildi"})


class StudentCharityModelView(APIView):
    def get(self, request, student_id):
        queryset = StudentPayment.objects.filter(student_id=student_id,
                                                 date__month__in=[9, 10, 11, 12, 1, 2, 3, 4, 5, 6], deleted=False,
                                                 status=True).annotate(month_number=ExtractMonth('date'))
        months_with_attendance = queryset.values_list('month_number', flat=True).distinct()
        all_months = [9, 10, 11, 12, 1, 2, 3, 4, 5, 6]
        missing_months = set(all_months) - set(months_with_attendance)
        month_names = [calendar.month_name[month] for month in sorted(missing_months)]
        return Response(month_names)

    def post(self, request, *args, **kwargs):
        data = json.loads(request.body)
        month = data.pop('date')
        month_number = list(calendar.month_name).index(month.capitalize())
        old_months = [9, 10, 11, 12]
        new_months = [1, 2, 3, 4, 5, 6, 7]
        current_year = int(data['year'])
        # if month_number in old_months:
        #     current_year -= 1
        # if month_number in new_months:
        #     current_year += 1
        date = datetime(year=current_year, month=int(datetime.now().month), day=int(datetime.now().day)).date()
        student_id = self.kwargs['student_id']
        student = get_object_or_404(Student, id=student_id)
        group = student.groups_student.first()
        attendance_per_month = AttendancePerMonth.objects.filter(student_id=student.id, month_date__month=month_number,
                                                                 month_date__year=current_year, group=group).first()
        if not attendance_per_month:
            return Response({"msg": "Shu yil shu oyga qarzdorlik yaratilmagan!"})

        if attendance_per_month.total_debt != attendance_per_month.payment and attendance_per_month.remaining_debt == 0:
            attendance_per_month.remaining_debt = attendance_per_month.total_debt
            attendance_per_month.save()

        payment_sum = int(request.data['payment_sum'])
        branch = request.data['branch']

        if attendance_per_month.remaining_debt >= payment_sum:
            attendance_per_month.remaining_debt -= payment_sum
            student_payment = StudentPayment.objects.create(student_id=student_id, payment_sum=payment_sum,
                                                            branch_id=branch, status=request.data['status'],
                                                            payment_type_id=request.data['payment_type'], date=date,
                                                            attendance=attendance_per_month,
                                                            reason=request.data['reason'])
            student_payment.save()
            discounts = StudentPayment.objects.filter(attendance=attendance_per_month, student_id=student_id,
                                                      branch_id=branch, deleted=False, status=True).all()
            student_payments = StudentPayment.objects.filter(attendance=attendance_per_month, student_id=student_id,

                                                             branch_id=branch,
                                                             deleted=False).all()
            total_payments = 0
            total_discount = 0
            for payment in discounts:
                total_discount += payment.payment_sum

            for payment in student_payments:
                total_payments += payment.payment_sum
            total_discount += payment_sum
            attendance_per_month.total_charity = total_discount
            attendance_per_month.payment = total_payments
            attendance_per_month.remaining_debt = attendance_per_month.total_debt - total_payments
            attendance_per_month.save()

            if attendance_per_month.remaining_debt == 0:
                attendance_per_month.status = True

        attendance_per_month.save()
        get_remaining_debt_for_student(student_id)

        return Response({"msg": "Chegirma muvaffaqiyatli yaratildi"})

    def put(self, request, *args, **kwargs):
        id = self.kwargs['student_id']
        sum = int(request.data.get('payment_sum'))
        payment = StudentPayment.objects.get(pk=id)
        attendance = AttendancePerMonth.objects.get(pk=payment.attendance.id)
        attendance.remaining_debt = attendance.total_debt - (
                attendance.payment - payment.payment_sum) + attendance.discount + sum
        attendance.payment = attendance.payment - payment.payment_sum + sum
        attendance.save()
        payment.payment_sum = sum
        payment.reason = request.data.get('reason', None)
        payment.save()
        get_remaining_debt_for_student(payment.student.id)
        return Response({"msg": "Chegirma muvaffaqiyatli o'zgartirildi"})


class GetYearView(APIView):
    def get(self, request):
        student_id = request.GET.get('student_id')
        if not student_id:
            return Response({"error": "student_id required"}, status=400)

        print("Student ID:", student_id)

        queryset = AttendancePerMonth.objects.filter(student_id=student_id)
        print("Attendance objects:", list(queryset.values('month_date')))

        years = queryset.annotate(year=ExtractYear('month_date')).values_list('year', flat=True).distinct()
        years = list(years)
        print("Extracted years:", years)

        current_year = datetime.today().year
        if current_year not in years:
            years.append(current_year)

        return Response({"years": sorted(years)})


class GetMonthView(APIView):
    def get(self, request):
        student_id = self.kwargs.get('student_id')
        year = self.kwargs.get('year')
        all_months = [9, 10, 11, 12, 1, 2, 3, 4, 5, 6]

        queryset = AttendancePerMonth.objects.filter(student_id=student_id, month_date__month__in=all_months,
                                                     month_date__year=year).annotate(
            month_number=ExtractMonth('month_date'))

        months_with_attendance = queryset.values_list('month_number', flat=True).distinct()

        missing_months = [month for month in all_months if month not in months_with_attendance]

        month_names = [calendar.month_name[month] for month in missing_months]
        return Response(month_names)

    def post(self, request):
        student_id = request.GET.get('student_id')
        year = request.GET.get('year')
        month_str = request.GET.get('month')

        if not (student_id and year and month_str):
            return Response({"error": "student_id, year, and month are required"}, status=400)

        try:
            month = datetime.strptime(month_str.strip().title(), '%B').month
            month_str_formatted = f"{month:02}"
        except ValueError:
            return Response({"error": f"Invalid month format: {month_str}"}, status=400)

        attendances = AttendancePerMonth.objects.filter(
            student_id=student_id,
            month_date__month=month_str_formatted,
            month_date__year=year
        ).first()

        if not attendances:
            return Response({"error": "Attendance not found"}, status=404)

        return Response({"sum": attendances.remaining_debt})


class GetStudentBalance(APIView):
    def get(self, request, user_id):
        student = Student.objects.get(user_id=user_id)
        balance = get_remaining_debt_for_student(student.id)
        return Response({"balance": balance}, status=status.HTTP_200_OK)


class MissingAttendanceListView2(generics.RetrieveAPIView):

    def get_queryset(self):
        student_id = self.kwargs.get('student_id')
        student = Student.objects.get(pk=student_id)
        group = student.groups_student.first()
        return AttendancePerMonth.objects.filter(student_id=student_id, group_id=group.id,
                                                 month_date__month__in=[9, 10, 11, 12, 1, 2, 3, 4, 5, 6]).annotate(
            month_number=ExtractMonth('month_date'))

    def retrieve(self, request, *args, **kwargs):
        queryset = self.get_queryset()
        student_id = self.kwargs.get('student_id')

        months_with_attendance = queryset.values_list('month_number', flat=True).distinct()
        data = []
        student = Student.objects.get(pk=student_id)

        group = student.groups_student.first()
        attendances = AttendancePerMonth.objects.filter(student_id=student_id, group_id=group.id).all().order_by(
            'month_date__year', 'month_date__month')

        for attendance in attendances:
            student_payemnt = StudentPayment.objects.filter(attendance=attendance, status=True

                                                            ).first()
            data.append({'id': attendance.id, 'month': attendance.month_date, 'total_debt': attendance.total_debt,
                         'remaining_debt': attendance.remaining_debt, 'payment': attendance.payment,
                         "discount": attendance.discount, "old_money": attendance.old_money,
                         "discount_sum": student_payemnt.payment_sum if student_payemnt else 0,
                         "discount_reason": student_payemnt.reason if student_payemnt else 0,
                         "discount_id": student_payemnt.id if student_payemnt else 0,
                         "reason": StudentCharity.objects.filter(
                             student_id=student_id).first().name if StudentCharity.objects.filter(
                             student_id=student_id).first() else None, 'cash':
                             attendance.studentpayment_set.filter(payment_type__name='cash', deleted=False,
                                                                  status=False,

                                                                  ).aggregate(total_sum=Sum('payment_sum'))[
                                 'total_sum'] or 0, 'bank':
                             attendance.studentpayment_set.filter(payment_type__name='bank', deleted=False,
                                                                  status=False,

                                                                  ).aggregate(total_sum=Sum('payment_sum'))[
                                 'total_sum'] or 0, 'click':
                             attendance.studentpayment_set.filter(payment_type__name='click', deleted=False,
                                                                  status=False,

                                                                  ).aggregate(total_sum=Sum('payment_sum'))[
                                 'total_sum'] or 0,

                         })

        all_months = [9, 10, 11, 12, 1, 2, 3, 4, 5, 6]
        missing_months = set(all_months) - set(months_with_attendance)
        month_names = [calendar.month_name[month] for month in sorted(missing_months)]

        return Response({"month": month_names, "data": data})


class StudentClassNumberUpdateView(APIView):
    def post(self, request):
        class_number = request.data.get('class_number')
        class_number_get = ClassNumber.objects.get(id=class_number)
        students_list = request.data.get('students')
        for st in students_list:
            student = Student.objects.get(id=st)
            student.class_number = class_number_get
            student.save()
        return Response({"msg": "O'quvchilar sinf raqami o'zgartirildi"}, status=status.HTTP_200_OK)


class ChangeDateDeletedStudent(APIView):
    def post(self, request):
        student_id = request.data.get('student_id')
        student = DeletedStudent.objects.get(id=student_id)
        student.deleted_date = request.data.get('del_date')
        student.save()
        return Response({"msg": "O'quvchi o'chirilgan vaqti o'zgartirildi"}, status=status.HTTP_200_OK)
